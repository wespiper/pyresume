"""
Enhanced DOCX text extraction using python-docx and lxml.
Handles headers, footers, text boxes, shapes, and all document parts.
"""
from typing import Optional, List, Set
import logging
import zipfile
import xml.etree.ElementTree as ET

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    Document = None

# Configure logging
logger = logging.getLogger(__name__)


class DOCXExtractor:
    """Extract text content from DOCX files with comprehensive coverage."""
    
    # XML namespaces used in DOCX files
    NAMESPACES = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'v': 'urn:schemas-microsoft-com:vml',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    }
    
    def __init__(self, debug: bool = False):
        if Document is None:
            raise ImportError("python-docx is required for DOCX extraction. Install with: pip install python-docx")
        self.debug = debug
        if debug:
            logging.basicConfig(level=logging.DEBUG)
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text content from a DOCX file including all document parts.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If DOCX cannot be read or processed
        """
        try:
            all_text_parts = []
            seen_text = set()  # To avoid duplicates
            
            # Method 1: Use python-docx for standard content
            doc = Document(file_path)
            
            # Extract headers
            for section in doc.sections:
                header = section.header
                if header:
                    header_text = self._extract_text_from_element(header._element)
                    if header_text and header_text not in seen_text:
                        all_text_parts.append(f"[Header] {header_text}")
                        seen_text.add(header_text)
                        if self.debug:
                            logger.debug(f"Found header: {header_text[:50]}...")
            
            # Extract main body paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip() and paragraph.text not in seen_text:
                    all_text_parts.append(paragraph.text)
                    seen_text.add(paragraph.text)
                    if self.debug:
                        logger.debug(f"Found paragraph: {paragraph.text[:50]}...")
            
            # Extract text from tables
            for table_idx, table in enumerate(doc.tables):
                table_texts = []
                for row_idx, row in enumerate(table.rows):
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text and cell_text not in seen_text:
                            row_text.append(cell_text)
                            seen_text.add(cell_text)
                    if row_text:
                        table_texts.append(' | '.join(row_text))
                
                if table_texts:
                    all_text_parts.append(f"\n[Table {table_idx + 1}]")
                    all_text_parts.extend(table_texts)
                    if self.debug:
                        logger.debug(f"Found table {table_idx + 1} with {len(table_texts)} rows")
            
            # Extract footers
            for section in doc.sections:
                footer = section.footer
                if footer:
                    footer_text = self._extract_text_from_element(footer._element)
                    if footer_text and footer_text not in seen_text:
                        all_text_parts.append(f"[Footer] {footer_text}")
                        seen_text.add(footer_text)
                        if self.debug:
                            logger.debug(f"Found footer: {footer_text[:50]}...")
            
            # Method 2: Direct XML parsing for text boxes, shapes, and other elements
            additional_text = self._extract_from_xml(file_path, seen_text)
            if additional_text:
                all_text_parts.extend(additional_text)
            
            # Join all text parts
            final_text = '\n'.join(all_text_parts)
            
            if self.debug:
                logger.debug(f"Total extracted text length: {len(final_text)} characters")
                logger.debug(f"Total unique text segments: {len(all_text_parts)}")
            
            return final_text
            
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {str(e)}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    def _extract_text_from_element(self, element) -> str:
        """Extract all text from an XML element recursively."""
        texts = []
        for elem in element.iter():
            if elem.tag.endswith('}t'):  # Text elements
                if elem.text:
                    texts.append(elem.text)
        return ' '.join(texts).strip()
    
    def _extract_from_xml(self, file_path: str, seen_text: Set[str]) -> List[str]:
        """
        Extract text from XML parts that python-docx might miss.
        This includes text boxes, shapes, SmartArt, etc.
        """
        additional_texts = []
        
        try:
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # List all XML files in the DOCX
                xml_files = [f for f in docx_zip.namelist() if f.endswith('.xml')]
                
                if self.debug:
                    logger.debug(f"Found {len(xml_files)} XML files in DOCX")
                
                # Process each XML file
                for xml_file in xml_files:
                    try:
                        content = docx_zip.read(xml_file)
                        root = ET.fromstring(content)
                        
                        # Extract text from various elements
                        texts = self._extract_texts_from_xml_root(root)
                        
                        for text in texts:
                            if text and text not in seen_text:
                                # Add context about where the text was found
                                if 'header' in xml_file:
                                    additional_texts.append(f"[Header XML] {text}")
                                elif 'footer' in xml_file:
                                    additional_texts.append(f"[Footer XML] {text}")
                                elif 'document' in xml_file:
                                    additional_texts.append(text)
                                else:
                                    additional_texts.append(f"[{xml_file}] {text}")
                                
                                seen_text.add(text)
                                
                                if self.debug:
                                    logger.debug(f"Found in {xml_file}: {text[:50]}...")
                    
                    except Exception as e:
                        if self.debug:
                            logger.debug(f"Could not parse {xml_file}: {e}")
                        continue
        
        except Exception as e:
            logger.warning(f"Could not perform deep XML extraction: {e}")
        
        return additional_texts
    
    def _extract_texts_from_xml_root(self, root) -> List[str]:
        """Extract all text content from an XML root element."""
        texts = []
        
        # Define all possible text-containing elements
        text_elements = [
            './/w:t',  # Regular text
            './/v:textpath',  # Text in VML shapes
            './/a:t',  # DrawingML text
            './/w:instrText',  # Field instruction text
            './/w:delText',  # Deleted text (track changes)
            './/w:noBreakHyphen',  # Non-breaking hyphens
        ]
        
        # Extract text from each type of element
        for xpath in text_elements:
            for elem in root.findall(xpath, self.NAMESPACES):
                if elem.text and elem.text.strip():
                    texts.append(elem.text.strip())
        
        # Also try without namespaces for compatibility
        for elem in root.iter():
            if elem.tag.endswith(('t', 'text', 'textpath')) and elem.text:
                text = elem.text.strip()
                if text and text not in texts:
                    texts.append(text)
        
        return texts
    
    def extract_structure(self, file_path: str) -> dict:
        """
        Extract document structure (headings, tables, etc.) for analysis.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing document structure information
        """
        try:
            doc = Document(file_path)
            structure = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'sections': len(doc.sections),
                'headers': 0,
                'footers': 0,
                'styles': set()
            }
            
            # Count headers and footers
            for section in doc.sections:
                if section.header:
                    structure['headers'] += 1
                if section.footer:
                    structure['footers'] += 1
            
            # Collect paragraph styles
            for paragraph in doc.paragraphs:
                if paragraph.style and paragraph.style.name:
                    structure['styles'].add(paragraph.style.name)
            
            structure['styles'] = list(structure['styles'])
            
            return structure
            
        except Exception as e:
            logger.error(f"Failed to extract structure: {e}")
            return {}